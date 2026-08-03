import io
import os
import re
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import Any, Optional

from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
import docx
from docx.oxml.ns import qn

# If modifying these scopes, delete the file token.json.
SCOPES = ["https://www.googleapis.com/auth/drive"]


class GoogleDocsService:
    def __init__(
        self, credentials_path: str = "config/secrets/credentials.json", token_path: str = "config/secrets/token.json"
    ):
        self.credentials_path = credentials_path
        self.token_path = token_path
        self.creds = None
        self.drive_service = None
        self._authenticate()

    def _authenticate(self):
        if os.path.exists(self.token_path):
            self.creds = Credentials.from_authorized_user_file(
                self.token_path, SCOPES
            )
        if not self.creds or not self.creds.valid:
            if self.creds and self.creds.expired and self.creds.refresh_token:
                self.creds.refresh(Request())
            else:
                if not os.path.exists(self.credentials_path):
                    raise ValueError(
                        f"Google OAuth credentials not found.\n\n"
                        "To fix this:\n"
                        "1. Go to https://console.cloud.google.com/\n"
                        "2. Create a project & enable the Google Docs API + Google Drive API\n"
                        "3. Create an OAuth client ID (Desktop app type)\n"
                        "4. Download the JSON and save it as credentials.json in config/secrets/\n"
                        f"   Expected location: {os.path.abspath(self.credentials_path)}"
                    )
                flow = InstalledAppFlow.from_client_secrets_file(
                    self.credentials_path, SCOPES
                )
                self.creds = flow.run_local_server(port=0)
            with open(self.token_path, "w") as token:
                token.write(self.creds.to_json())

        self.drive_service = build("drive", "v3", credentials=self.creds)

    def download_as_docx(self, doc_id: str, output_path: str) -> str:
        """Download a Google Doc as a .docx file."""
        request = self.drive_service.files().export_media(
            fileId=doc_id,
            mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()

        with open(output_path, "wb") as f:
            f.write(fh.getvalue())
        return output_path

    def upload_docx(
        self, docx_path: str, filename: str, parents: Optional[list] = None
    ) -> str:
        """Upload a .docx file to Google Drive and convert to Google Doc. Returns the new Doc ID."""
        file_metadata = {
            "name": filename,
            "mimeType": "application/vnd.google-apps.document",
        }
        if parents:
            file_metadata["parents"] = parents

        media = MediaFileUpload(
            docx_path,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            resumable=True,
        )
        file = (
            self.drive_service.files()
            .create(body=file_metadata, media_body=media, fields="id")
            .execute()
        )
        return file.get("id")

    # ── Formatting-preserving docx manipulation ──────────────────────────

    @staticmethod
    def parse_docx_structure(docx_path: str) -> tuple:
        """Parse a .docx file into sections grouped by headings.

        Returns (doc, sections) where doc is the python-docx Document object
        (needed for saving later) and sections is a list of dicts:
            [{ "heading": "Profile", "heading_para": <element>,
               "paragraphs": [{"text": "...", "element": <elem>, "font": {...}}, ...] }, ...]

        Headings are detected by paragraph style name ("Heading N") or by
        heuristics (bold + larger font than body).
        """
        doc = docx.Document(docx_path)
        sections: list[dict] = []
        current_section: Optional[dict] = None

        # First pass: determine the "body" font size (most common size)
        body_font_size = GoogleDocsService._detect_body_font_size(doc)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_heading = GoogleDocsService._is_heading_paragraph(
                para, body_font_size
            )

            if is_heading:
                font_info = GoogleDocsService._capture_first_run_font(para)
                current_section = {
                    "heading": text,
                    "heading_para": para._element,
                    "heading_font": font_info,
                    "paragraphs": [],
                }
                sections.append(current_section)
            elif current_section is not None:
                font_info = GoogleDocsService._capture_first_run_font(para)
                current_section["paragraphs"].append(
                    {"text": text, "element": para._element, "font": font_info}
                )
            else:
                # Body text before any heading — create a preamble section
                font_info = GoogleDocsService._capture_first_run_font(para)
                current_section = {
                    "heading": "",
                    "heading_para": None,
                    "heading_font": None,
                    "paragraphs": [],
                }
                sections.append(current_section)
                current_section["paragraphs"].append(
                    {"text": text, "element": para._element, "font": font_info}
                )

        return doc, sections

    @staticmethod
    def _detect_body_font_size(doc: docx.Document) -> float:
        """Find the most common font size among non-bold, non-heading paragraphs."""
        from collections import Counter

        sizes: list[float] = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style_name = (para.style.name or "").lower()
            if "heading" in style_name:
                continue
            for run in para.runs:
                if run.font.size and not run.bold:
                    sizes.append(run.font.size / 12700)  # EMU to pt
                    break
        if not sizes:
            # Fallback: include bold runs too
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                for run in para.runs:
                    if run.font.size:
                        sizes.append(run.font.size / 12700)
                        break
        if not sizes:
            return 11.0
        size_counts = Counter(round(s, 1) for s in sizes)
        return size_counts.most_common(1)[0][0]

    @staticmethod
    def _is_heading_paragraph(para, body_font_size: float) -> bool:
        """Determine if a paragraph is a heading."""
        style_name = (para.style.name or "").lower()

        # Explicit heading styles
        if "heading" in style_name or style_name.startswith("toc"):
            return True

        # Heuristic: short text, bold, larger than body font
        text = para.text.strip()
        if len(text) > 150:
            return False

        first_run = para.runs[0] if para.runs else None
        if first_run is None:
            return False

        is_bold = first_run.bold
        font_size = (first_run.font.size or 0) / 12700  # EMU to pt

        # Heading if: bold AND font size > body size by at least 1pt
        if is_bold and font_size > body_font_size + 0.5:
            return True

        # Also: all-caps short text with no trailing punctuation
        if text.isupper() and len(text) < 80 and not text.endswith("."):
            return True

        return False

    @staticmethod
    def _capture_first_run_font(para) -> Optional[dict]:
        """Capture font properties from the first run of a paragraph."""
        if not para.runs:
            return None
        run = para.runs[0]
        font = run.font
        return {
            "name": font.name,
            "size": font.size,
            "bold": font.bold,
            "italic": font.italic,
            "underline": font.underline,
            "color": font.color.rgb if font.color and font.color.rgb else None,
        }

    @staticmethod
    def extract_section_text(sections: list[dict]) -> str:
        """Join section data into a single string with section markers for the LLM."""
        parts: list[str] = []
        for sec in sections:
            heading = sec.get("heading", "")
            if heading:
                parts.append(f"## {heading}")
            for p in sec.get("paragraphs", []):
                parts.append(p["text"])
            parts.append("")  # blank line between sections
        return "\n".join(parts)

    @staticmethod
    def apply_tailored_sections(
        doc: docx.Document, sections: list[dict], tailored_text: str, output_path: str
    ) -> None:
        """Apply the LLM's tailored text back to the docx, preserving formatting.

        Parses the tailored markdown into sections, matches them to the original
        docx sections, and replaces paragraph text in-place while keeping
        first-run formatting (font, size, bold, italic, color).
        """
        tailored_sections = GoogleDocsService._parse_markdown_sections(
            tailored_text
        )

        for orig_sec in sections:
            heading = orig_sec.get("heading", "")
            matched = GoogleDocsService._match_section(
                heading, tailored_sections
            )

            if matched is None:
                continue

            orig_paras = orig_sec.get("paragraphs", [])
            new_paras = matched.get("paragraphs", [])

            # Replace text in-place for matching paragraphs
            for i, orig_p in enumerate(orig_paras):
                if i < len(new_paras):
                    GoogleDocsService._replace_paragraph_text(
                        orig_p["element"], new_paras[i], orig_p.get("font")
                    )
                else:
                    # Fewer new paragraphs — clear the extras
                    GoogleDocsService._replace_paragraph_text(
                        orig_p["element"], "", orig_p.get("font")
                    )

            # More new paragraphs than original — clone the last paragraph
            if len(new_paras) > len(orig_paras) and orig_paras:
                last_elem = orig_paras[-1]["element"]
                last_font = orig_paras[-1].get("font")
                parent = last_elem.getparent()
                insert_after = last_elem
                for extra_text in new_paras[len(orig_paras) :]:
                    cloned = deepcopy(last_elem)
                    # Clear paragraph text in the clone
                    for r in cloned.findall(qn("w:r")):
                        cloned.remove(r)
                    # Find the insertion index and add after the reference
                    parent_elem = list(parent)
                    try:
                        idx = parent_elem.index(insert_after)
                        parent.insert(idx + 1, cloned)
                    except ValueError:
                        parent.append(cloned)
                    GoogleDocsService._replace_paragraph_text(
                        cloned, extra_text, last_font
                    )
                    insert_after = cloned

            # Also update the heading text if it changed
            if heading and matched.get("heading") and orig_sec.get("heading_para") is not None:
                new_heading = matched["heading"]
                if new_heading != heading:
                    GoogleDocsService._replace_paragraph_text(
                        orig_sec["heading_para"],
                        new_heading,
                        orig_sec.get("heading_font"),
                    )

        doc.save(output_path)

    @staticmethod
    def _replace_paragraph_text(
        para_element, new_text: str, font_info: Optional[dict]
    ) -> None:
        """Replace all text in a paragraph element, preserving formatting.

        Clears all existing runs, then adds a single new run with the replacement
        text, applying font properties from font_info.
        """
        # Remove all existing runs
        for run_elem in para_element.findall(qn("w:r")):
            para_element.remove(run_elem)

        # Create a new run with the replacement text
        new_run = docx.oxml.OxmlElement("w:r")
        rpr = docx.oxml.OxmlElement("w:rPr")

        if font_info:
            if font_info.get("name"):
                rfont = docx.oxml.OxmlElement("w:rFonts")
                rfont.set(qn("w:ascii"), font_info["name"])
                rfont.set(qn("w:hAnsi"), font_info["name"])
                rpr.append(rfont)
            if font_info.get("size"):
                sz = docx.oxml.OxmlElement("w:sz")
                # font.size is in EMU; w:sz attribute is in half-points
                # 1 pt = 12700 EMU = 2 half-pts, so half-pts = EMU / 6350
                sz.set(qn("w:val"), str(int(font_info["size"] / 6350)))
                rpr.append(sz)
            if font_info.get("bold"):
                b = docx.oxml.OxmlElement("w:b")
                rpr.append(b)
            if font_info.get("italic"):
                i = docx.oxml.OxmlElement("w:i")
                rpr.append(i)
            if font_info.get("underline"):
                u = docx.oxml.OxmlElement("w:u")
                u.set(qn("w:val"), "single")
                rpr.append(u)
            if font_info.get("color"):
                color = docx.oxml.OxmlElement("w:color")
                color.set(qn("w:val"), str(font_info["color"]))
                rpr.append(color)

        if len(rpr) > 0:
            new_run.append(rpr)

        text_elem = docx.oxml.OxmlElement("w:t")
        text_elem.text = new_text
        text_elem.set(qn("xml:space"), "preserve")
        new_run.append(text_elem)
        para_element.append(new_run)

    @staticmethod
    def _parse_markdown_sections(text: str) -> list[dict]:
        """Parse tailored markdown into sections matching the docx structure.

        Splits on `## Heading` markers. Lines before the first heading become
        a preamble section with empty heading. Body lines within each section
        become paragraphs (blank lines between paragraphs are merged).
        """
        sections: list[dict] = []
        lines = text.split("\n")

        current_heading = ""
        current_paras: list[str] = []

        for line in lines:
            stripped = line.strip()
            # Detect ## section headings
            if re.match(r"^#{1,3}\s+", stripped):
                if current_paras or current_heading:
                    sections.append(
                        {"heading": current_heading, "paragraphs": current_paras}
                    )
                current_heading = re.sub(r"^#{1,3}\s+", "", stripped).strip()
                current_paras = []
            elif stripped == "":
                continue
            elif stripped == "---":
                break
            else:
                current_paras.append(stripped)
        if current_paras or current_heading:
            sections.append(
                {"heading": current_heading, "paragraphs": current_paras}
            )

        return sections

    @staticmethod
    def _match_section(
        heading: str, tailored_sections: list[dict]
    ) -> Optional[dict]:
        """Fuzzy-match a docx section heading to a tailored markdown section."""
        if not heading:
            # Preamble section: return the first section with an empty heading
            for sec in tailored_sections:
                if not sec.get("heading"):
                    return sec
            return None

        heading_lower = heading.lower().strip()

        # Exact match first
        for sec in tailored_sections:
            if sec.get("heading", "").lower().strip() == heading_lower:
                return sec

        # Substring match (e.g., "Work Experience" matches "## Work Experience")
        for sec in tailored_sections:
            sec_heading = sec.get("heading", "").lower().strip()
            if heading_lower in sec_heading or sec_heading in heading_lower:
                return sec

        # Word overlap match (e.g., "Skills" matches "Key Skills")
        heading_words = set(heading_lower.split())
        for sec in tailored_sections:
            sec_words = set(sec.get("heading", "").lower().split())
            if heading_words & sec_words:
                return sec

        return None
